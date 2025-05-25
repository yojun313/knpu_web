import sys
import json
from openai import OpenAI
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from doc_process import doc_process
from docx.shared import RGBColor
import os
from doc_process import doc_process
from dotenv import load_dotenv

load_dotenv()

input_word = sys.argv[1]

class statement_gpt:
    def __init__(self, input_word):
        
        self.statement_default_document = Document(os.path.join(os.path.dirname(__file__), "진술조서 양식.docx")) # 진술조서 양식 문서
        self.docx_statement_storage = os.path.join(os.path.dirname(__file__), "docx_statement_storage")
        self.storage_json = os.path.join(os.path.dirname(__file__), "public/storage.json")
        self.doc_process = doc_process()
        self.input_dic = os.path.join(os.path.dirname(__file__), "public/input_dic.txt")
        self.input_word = input_word
        
        with open(self.input_dic, "r", encoding="utf-8") as f:
            lines = f.readlines()
        for i in range(len(lines)):
            if lines[i].replace('\n', '') == input_word:
                data_str = lines[i+1].replace("'", "\"")
                self.form_data = json.loads(data_str)
                break

    def main(self):
        
        statement_query = "GPT4를 사용해서 진술 조서를 작성할거야. 다음에 내가 제시할 '사용자의 입력값'을 바탕으로 진술 조서를 작성해줘. 단, 다음 조건을 만족시켜줘."\
                "진술 조서와 관련없는 내용을 빼고 출력할 것(예를 들어, '당연히 가능합니다', '위 진술 조서는 가상의 정보를 바탕으로 작성된 것이며, 실제 고소장 작성시에는 법률 전문가의 도움을 받는 것을 권장드립니다.' 같은 말)"\
                "고소장 딕셔너리에 담겨야 하는 '입력값'은 아래와 같아."+ str(self.form_data) + "."\
                "나는 너가 준 답변을 바로 파이썬 코드에 딕셔너리 코드에 넣을거야. 그러므로 너는 내가 줄 이 형식에 맞게 파이썬 딕셔너리 형태로 '출력값'을 출력해줘."\
                "지금 준 '입력값'의 key를 전부 구어체(존댓말로)로 바꿔. 예를 들어 고소인 성명은 고소인의 성명은 무엇입니까?로, 전부 물어보는 형식으로 바꿔."\
                "그리고 바꾼 구어체의 key를 output의 key로 넣어줘. 근데 조건이 '문:'으로 시작해야해. 예를 들어서 입력값의 key가 이름이라면 '출력값'의 형식은 다음과 같아."\
                "출력값: key: '문: 고소인의 이름은 무엇입니까' value: '답: ~~~입니다'"\
                "'출력값'의 value에는 입력값의 value를 존댓말로 바꿔서 적으면 돼. 예를 들어 '입력값'의 value에 이름이 '홍길동'이라고 쓰여있으면 출력값의 value에는 '답: 홍길동입니다' 이렇게 적으면 돼."\
                "'출력값'의 value에는 답을 자연스럽게 출력해야돼. 예를 들어 '있음'은 있습니다로, '합의하였음'은 합의하였습니다로, '아니요'는 아닙니다로. 이렇게 '입니다'만 붙이는 것이 아닌 자연스럽게 처리해줘"\
                "이런 방식으로 출력값(딕셔너리)의 key에는 문: 으로 시작하는 존댓말 질문을 적으면 되고 출력값(딕셔너리)의 value에는 답: 으로 시작하는 존댓말 답변을 적으면 돼."\
                "너는 다른 말 없이 완성된 딕셔너리만 출력해."\
                "너가 출력해준 답변을 바로 json.loads()함수에 넣을거니깐 에러 안나게 딕셔너리를 한줄로 출력해줘"\
                "너가 결과값으로 나에게 주는 딕셔너리는 key와 value모두 꼭 쌍따옴표("")로 이루어져있어야해"\
                "꼭 한줄로, 변수 빼고 꼭 {으로 시작해서 }으로 끝나게 답변 줘봐. 딕셔너리 말고 아무것도 답변에 추가하지마."
        
        statement_content = self.gpt_generate(statement_query)
        statement_data = json.loads(statement_content)
        self.statement_generate(self.form_data, statement_data)
        self.drive_upload(self.docx_statement_word_name, self.docx_statement_word_path)
        
        
    def gpt_generate(self, query):
        client = OpenAI(api_key = os.getenv("OPENAI_API_KEY"))
        model = "gpt-4"
        response = client.chat.completions.create(
            model = model,
            messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": query},
                ]
        )
        content = response.choices[0].message.content
        return content
    
    def add_sentence(self, target_sentence, new_text):
    # 문서의 모든 단락을 순회하며 특정 문장을 찾음
        for i, paragraph in enumerate(self.sue_default_document.paragraphs):
            if target_sentence in paragraph.text:
                # 새로운 단락을 추가하고 이를 현재 단락 뒤로 이동
                new_paragraph = self.sue_default_document.add_paragraph(new_text)
                p = self.sue_default_document.paragraphs[i]._element
                p.addnext(new_paragraph._element)
                break  # 첫 번째 일치하는 문장을 찾으면 반복 중단
    
    def statement_generate(self, form_data, statement_data):
        statement_doc = Document()
        
        title = statement_doc.add_heading('진술 조서', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = "Arial"
        title.runs[0].font.color.rgb = RGBColor(0,0,0)
        title.bold = True
        
        statement_list_q = list(statement_data.keys())
        statement_list_a = list(statement_data.values())
        
        statement_doc.add_paragraph("성                  명:  "+form_data['고소인'])
        statement_doc.add_paragraph("주민등록번호:  "+str(form_data['고소인 주민등록번호']))
        statement_doc.add_paragraph("직                  업:  "+form_data['고소인 직업'])
        statement_doc.add_paragraph("주                  거:  "+form_data['고소인 주소'])
        statement_doc.add_paragraph("휴대         전화:  "+form_data['고소인 전화'])
        statement_doc.add_paragraph("전자         우편:  "+form_data['고소인 이메일'])
        statement_doc.add_paragraph("\n")
        statement_doc.add_paragraph("위의 사람은 피의자 {}에 대한 사기 피의사건에 관하여 {}에 {} 조사실에 임의 출석하여 다음과 같이 진술하다.".format(form_data['피고소인'], form_data['고소일자'], form_data['제출 경찰서']))
        statement_doc.add_paragraph("\n\n")
        
        statement_doc.add_paragraph("1. 피의자와의 관계").bold = True

        statement_doc.add_paragraph("   "+statement_list_a[13].replace("답: ", ""))
        statement_doc.add_paragraph("\n")
        
        statement_doc.add_paragraph("2. 피의사실과의 관계").bold = True
        statement_doc.add_paragraph("   저는 피의 사실과 관련하여 고소인 자격으로서 출석하였습니다.")
        statement_doc.add_paragraph("\n 이 때 진술의 취지를 더욱 명백히 하기 위하여 다음과 같이 임의로 문답하다.\n\n")
        
        for i in range(12, len(statement_list_a)-2):
            statement_doc.add_paragraph(statement_list_q[i]).bold = True
            statement_doc.add_paragraph(str(statement_list_a[i]))
            statement_doc.add_paragraph("\n")
        
        statement_doc.add_paragraph("\n\n")
        
        table1 = statement_doc.add_table(rows=1, cols=1)
        table1.style = 'Table Grid'
        
        cell1 = table1.cell(0, 0)  # 첫 번째 행, 첫 번째 열의 셀 선택
        paragraph1 = cell1.paragraphs[0]
        paragraph1.text = "\n\n위의 조서를 진술자에게 열람하게 하였던 바 진술한 대로 오기나 증감·변경할 것이 없다고 말하므로 서명(기명날인)하게 하다.\
                         \n\n\n\n\n\n진술자  {} (인)\n\n\n\n{}\n\n\n\n\n".format(form_data['고소인'], form_data['고소일자'])
        paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        checkdoc = statement_doc.add_paragraph("\n\n수사 과정 확인서")
        checkdoc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        checkdoc.runs[0].font.name = "Arial"
        checkdoc.runs[0].font.color.rgb = RGBColor(0,0,0)
        checkdoc.bold = True
        
        table2 = statement_doc.add_table(rows=7, cols=2)
        table2.style = 'Table Grid'
        last_row = table2.rows[6]  # 두 번째 행 선택
        last_row.cells[0].merge(last_row.cells[1])  # 여섯 번째 및 일곱 번째 셀 병합
        
        cell2 = table2.cell(0,0)
        paragraph2 = cell2.paragraphs[0]
        paragraph2.text = "구분"
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cell3 = table2.cell(0,1)
        paragraph3 = cell3.paragraphs[0]
        paragraph3.text = "내용"
        paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cell4 = table2.cell(1,0)
        cell4.text = "1. 조사 장소 도착시각"
        
        cell5 = table2.cell(2,0)
        cell5.text = "2. 조사 시작시각 및 종료시각"
        
        cell6 = table2.cell(3,0)
        cell6.text = "3. 조서 열람 시작시각 및 종료시각"
        
        cell7 = table2.cell(4,0)
        cell7.text = "4. 기타 조사과정 진행경과 확인에 필요한 사항"
        
        cell8 = table2.cell(5,0)
        cell8.text = "5. 조사과정 기재사항에 대한 이의 제기나 의견진술 여부 및 그 내용"
        
        cell9 = table2.cell(6,0)
        paragraph3 = cell9.paragraphs[0]
        paragraph3.text = "\n\n{}\n사법경찰관 경감 ___은 {}을(를) 조사한 후, 위와 같은 사항에 대해 {}으로부터 확인받음\n\n\n확인자 {} (인)\n\n사법경찰관 ___ (인)\n\n".format(form_data['고소일자'], form_data['고소인'], form_data['고소인'], form_data['고소인'])
        paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.docx_statement_word_name = self.input_word.replace("고소장", "진술 조서")
        self.docx_statement_word_path = os.path.join(self.docx_statement_storage, self.input_word.replace("고소장", "진술 조서"))
        statement_doc.save(self.docx_statement_word_path)

    def drive_upload(self, word_name, word_path):
        pdf_name = word_name
        pdf_path = self.doc_process.convert_statement_word_to_pdf(word_name, word_path) # 워드 문서 pdf로 변환한 뒤, 로컬 pdf 경로 획득
        
        edit_link = self.doc_process.word_upload(word_name, word_path) # 워드 드라이브에 업로드 한 뒤 편집 링크 획득
        pdf_id = self.doc_process.pdf_upload(pdf_name, pdf_path) # pdf 드라이브에 업로드 한 뒤 pdf_id 획득
        
        with open(self.input_dic, "a", encoding="utf-8") as file:
            # 텍스트 추가
            file.write(word_name+"\n")
            file.write(str(self.form_data)+"\n")

        print(edit_link)
        print(pdf_id)
        print(word_name)

sue_gpt = statement_gpt(input_word)
sue_gpt.main()

