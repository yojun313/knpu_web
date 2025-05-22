import os
from docx import Document
from doc_process import doc_process
import json

# 워드 파일이 있는 폴더 경로
word_folder_path = os.path.join(os.path.dirname(__file__), 'docx_storage')
pdf_folder_path = os.path.join(os.path.dirname(__file__), 'pdf_storage')

doc_process = doc_process()

# 폴더 내 모든 파일 리스트를 가져옴
file_list = os.listdir(word_folder_path)
file_list = sorted(file_list, key=lambda x: os.path.getctime(os.path.join(word_folder_path, x)))

# .docx 파일들만 필터링
docx_files_name = [file for file in file_list if file.endswith('.docx')]
sue_info = []

num = 1
for docx_file_name in docx_files_name:
    word_file_path = os.path.join(word_folder_path, docx_file_name)  # 워드 파일 경로
    pdf_file_name = docx_file_name.replace("docx", 'pdf') # pdf 파일 이름
    pdf_file_path = os.path.join(pdf_folder_path, pdf_file_name)
    writer = word_file_path.split("_")[2]
    pdf_id = doc_process.pdf_upload(pdf_file_name, pdf_file_path)
    
    doc = Document(word_file_path)
    for paragraph in doc.paragraphs:
        if "월" in paragraph.text and "일" in paragraph.text:
            sue_date = paragraph.text.replace("\n                                                 ", "")
    receiver = doc.tables[1].rows[0].cells[1].paragraphs[0].text
    sin = docx_file_name.split("_")[2]
    formatted_number = f'{num:06d}'
    formatted_number = f'2024-{formatted_number}'
 
    sue_info.append([num, formatted_number, sin, writer, receiver, sue_date, 'https://drive.google.com/file/d/'+pdf_id +'/preview' , docx_file_name])
    num += 1

print(json.dumps(sue_info))

