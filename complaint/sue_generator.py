# -*- coding: utf-8 -*-
import sys
import json
from openai import OpenAI
from docx import Document
from datetime import datetime
import os
from doc_process import doc_process
from dotenv import load_dotenv

load_dotenv()

first_form_data = json.loads(sys.argv[1])
second_form_data = json.loads(sys.argv[2])

form_data = {**first_form_data, **second_form_data}
for key, value in form_data.items():
    form_data[key] = form_data[key].replace("\t", " ").replace("\r", " ").replace("\n", " ")


class sue_gpt:
    def __init__(self, form_data):
        self.form_data = form_data # 웹사이트로부터 입력받은 데이터 
        self.output_data = {
            "고소인 성명": "(고소인의 이름을 입력해줘)",
            "고소인 주민등록번호": "(고소인의 주민등록번호를 입력해줘)",
            "고소인 주소": "(고소인의 주소를 입력해줘)",
            "고소인 직업": "(고소인의 직업을 입력해줘)",
            "고소인 전화": "(고소인의 전화번호를 입력해줘)",
            "고소인 이메일": "(고소인의 이메일을 입력해줘)",
            "피고소인 성명": "(피고소인의 이름을 입력해줘. 없을 경우 공백으로 놓으면 돼)",
            "피고소인 주민등록번호": "(피고소인의 주민등록번호를 입력해줘. 없을 경우 공백으로 놓으면 돼)",
            "피고소인 주소": "(피고소인의 주소를 입력해줘. 없을 경우 공백으로 놓으면 돼)",
            "피고소인 직업": "(피고소인의 직업을 입력해줘. 없을 경우 공백으로 놓으면 돼)",
            "피고소인 전화": "(피고소인의 전화번호를 입력해줘. 없을 경우 공백으로 놓으면 돼)",
            "피고소인 이메일": "(피고소인의 이메일을 입력해줘. 없을 경우 공백으로 놓으면 돼)",
            "피고소인 기타사항": "(고소인과의 관계 및 피고소인의 인적사항과 연락처를 정확히 알 수 없을 경우 피고소인의 성별, 특징적 외모, 인상착의 등을 입력해줘. 기타 피고소인을 특정할 수 있는 내용을 입력하면돼)",
            "고소 취지": "(죄명 및 피고소인에 대한 처벌 의사를 포함해야해. 형식은 '고소인은 피고소인을 ~~죄로 고소하오니 처벌하여 주시기 바랍니다.'이거야.)",
            "범죄 사실": "(범죄사실은 형법 등 처벌법규에 해당하는 사실에 대하여 일시, 장소, 범행방법, 결과 등을 구체적으로 특정하여 기재해야 하며, 고소인이 알고 있는 지식과 경험, 증거에 의해 사실로 인정되는 내용을 기재해줘.\
                         (예시: 피고소인는 무직으로 별다른 수입이 없고, 20,000,000원 상당의 채무를 부담하고 있었으며, 별다른 재산이 없어 피해자로부터 돈을 빌리더라도 갚을 의사나 능력이 없습니다.\
                          그럼에도 불구하고 피고소인은 2021. 0. 0. 00:00경 ○○구 ○○로에 있는 고소인의 집에서, 고소인에게 “10,000,000원만 빌려 주면 월 3%의 이자를 지급하고, 2개월 후에 틀림없이 갚겠다”고 거짓말하였습니다.\
                          피고소인은 이와 같이 고소인을 기망하여 이에 속은 고소인으로부터 즉석에서 차용금 명목으로 10,000,000원을 교부받았습니다.)이 예시를 보고 학습해서 범죄 사실을 적으면 돼.\
                          문장은 경어체(존댓말)로 끝나야하고 자신을 가리킬 때는 '고소인'이라고 지칭해야돼. 그리고 피해 일시에서 날짜 뿐만 아니라 시간도 꼭 포함해야돼)",
            "고소 이유": "(고소이유에는 피고소인의 범행 경위 및 정황, 고소를 하게 된 동기와 사유 등 범죄사실을 뒷받침하는 내용을 간략, 명료하게 기재해야 해. 그리고 합의 여부, 처벌 의사에 따라 '범죄사실과 같이 피고소인은 사기범행을 \
                          하였습니다. 피고소인과는 합의하지 않았으며/합의하였으며, 피고소인의 처벌을 원합니다/원하지 않습니다'의 문장이 포함되어 있어야해.)",
            "증거 자료": "(증거자료를 입력해줘. 없으면 공백으로 놓으면 돼)",
            "중복 고소 여부": "(있으면 '있음', 없으면 '없음'이라고 입력해줘. 물론 따옴표는 뺴야돼. 즉 있음 또는 없음으로)",
            "관련 형사사건 수사 유무": "(있으면 '있음', 없으면 '없음'이라고 입력해줘. 물론 따옴표는 뺴야돼. 즉 있음 또는 없음으로)",
            "기타": "(너가 위에서 말하지 못한 것들을 여기에 입력해줘. 없으면 공백으로 놓으면 돼)",
            "고소일": "('~~~~년 ~~월 ~~일' 로 입력해줘. (예: 2024년 1월 20일) 예를 들면 이런 형식으로. 물론 따옴표 없이, 01월 11일 이런식으로 월의 숫자가 한자리라고 해서 앞에 0을 붙이지 마)",
            "제출 경찰서": "('~~경찰서 귀중' 으로 입력해줘. 물론 따옴표 없이)"
        }
        self.document = Document(os.path.join(os.path.dirname(__file__), "public/고소장 표준 양식(경찰청).docx")) # 고소장 양식 문서
        self.docx_storage = os.path.join(os.path.dirname(__file__), "docx_storage")
        self.api_key = os.getenv("OPENAI_API_KEY")
        self.storage_json = os.path.join(os.path.dirname(__file__), "public/storage.json")
        self.doc_process = doc_process()
        self.input_dic = os.path.join(os.path.dirname(__file__), "public/input_dic.txt")
    
    def main(self):
        query = "GPT4를 사용해서 고소장을 작성할거야. 다음에 내가 제시할 '고소 내용'을 바탕으로 고소장을 작성해줘. 단, 다음 조건을 만족시켜줘."\
                "고소장과 관련없는 내용을 빼고 출력할 것(예를 들어, '당연히 가능합니다', '위 고소장은 가상의 정보를 바탕으로 작성된 것이며, 실제 고소장 작성시에는 법률 전문가의 도움을 받는 것을 권장드립니다.' 같은 말)"\
                "고소장 딕셔너리에 담겨야 하는 '입력값'은 아래와 같아."+ str(self.form_data) + "."\
                "나는 너가 준 답변을 바로 파이썬 코드에 딕셔너리 코드에 넣을거야. 그러므로 너는 내가 줄 이 형식에 맞게 파이썬 딕셔너리 형태로 답을 출력해줘."\
                "지금 여기서부턴 딕셔너리야." + str(self.output_data) + "."\
                "딕셔너리의 key값은 절대 건드리지 말고 value내에는 괄호 안에 내가 너한테 할 명령을 적어놨어."\
                "이 명령에 따라서 너는 괄호와 괄호 안의 내용을 지우고 너의 답변을 딕셔너리의 value부분에 넣으면 돼."\
                "너는 다른 말 없이 완성된 딕셔너리만 출력해."\
                "너가 출력해준 답변을 바로 json.loads()함수에 넣을거니깐 에러 안나게 딕셔너리를 한줄로 출력해줘"\
                "너가 결과값으로 나에게 주는 딕셔너리는 key와 value모두 꼭 쌍따옴표("")로 이루어져있어야해"\
                "꼭 한줄로, 변수 빼고 꼭 {으로 시작해서 }으로 끝나게 답변 줘봐. 딕셔너리 말고 아무것도 답변에 추가하지마."
        content = self.gpt_generate(query)
        data = json.loads(content)

        word_name, word_path = self.word_generate(data)
        
        self.drive_upload(word_name, word_path)
        
    def gpt_generate(self, query):
        client = OpenAI(api_key = self.api_key)
        model = "gpt-4"
        response = client.chat.completions.create(
            model = model,
            messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": query},
                ]
        )

        content = response.choices[0].message.content
        return content
    
    def add_sentence(self, target_sentence, new_text):
    # 문서의 모든 단락을 순회하며 특정 문장을 찾음
        for i, paragraph in enumerate(self.document.paragraphs):
            if target_sentence in paragraph.text:
                # 새로운 단락을 추가하고 이를 현재 단락 뒤로 이동
                new_paragraph = self.document.add_paragraph(new_text)
                p = self.document.paragraphs[i]._element
                p.addnext(new_paragraph._element)
                break  # 첫 번째 일치하는 문장을 찾으면 반복 중단
    
    def word_generate(self, data):
        
        now = datetime.now()
        
        table_requester = self.document.tables[0]
        for row in table_requester.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text == "고소인 성명":
                        para.text = data["고소인 성명"]
                    elif para.text == "고소인 주민등록번호":
                        para.text = data["고소인 주민등록번호"]
                    elif para.text == "고소인 주소":
                        para.text = data["고소인 주소"]
                    elif para.text == "고소인 직업":
                        para.text = data["고소인 직업"]
                    elif para.text == "고소인 전화":
                        para.text = data["고소인 전화"]
                    elif para.text == "고소인 이메일":
                        para.text = data["고소인 이메일"]
        table_receiver = self.document.tables[1]

        for row in table_receiver.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text == "피고소인 성명":
                        para.text = data["피고소인 성명"]
                    elif para.text == "피고소인 주민등록번호":
                        para.text = data["피고소인 주민등록번호"]
                    elif para.text == "피고소인 주소":
                        para.text = data["피고소인 주소"]
                    elif para.text == "피고소인 직업":
                        para.text = data["피고소인 직업"]
                    elif para.text == "피고소인 전화":
                        para.text = data["피고소인 전화"]
                    elif para.text == "피고소인 이메일":
                        para.text = data["피고소인 이메일"]
                    elif para.text == "피고소인 기타사항":
                        para.text = data["피고소인 기타사항"]

        self.add_sentence("(죄명 및 피고소인에 대한 처벌의사 기재)", "\n"+data["고소 취지"]) # \n 무조건 추가 필요
        self.add_sentence("4. 범죄사실*", "\n"+data["범죄 사실"]) # \n 무조건 추가 필요
        self.add_sentence("5. 고소이유", "\n"+data["고소 이유"]) # \n 무조건 추가 필요
        self.add_sentence("6. 증거자료", "\n"+data["증거 자료"]) # \n 무조건 추가 필요

        checklist_table = self.document.tables[2]

        for row in checklist_table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text == "본 고소장과 같은 내용의 고소장을 다른 검찰청 또는 경찰서에 제출하거나 제출하였던 사실이 ":
                        para.add_run(data["중복 고소 여부"])
                    elif para.text == "본 고소장에 기재된 범죄사실과 관련된 사건 또는 공범에 대하여 검찰청이나 경찰서에서 수사 중에 ":
                        para.add_run(data["관련 형사사건 수사 유무"])

        self.add_sentence("8. 기타", "\n"+data["기타"]) # \n 무조건 추가 필요
        self.add_sentence("무고죄로 처벌받을 것임을 서약합니다.", "\n                                                 "+data["고소일"])
        self.add_sentence("고소대리의 경우에는 제출인을 기재하여야 합니다.", "\n\n                                                   "+data["제출 경찰서"])
        suedocument_name = "GPT 고소장_"+str(data["고소인 성명"])+"_"+str(self.form_data["고소 죄명"])+"_"+str(now.strftime("%H%M%S"))+".docx"
        suedocument = os.path.join(self.docx_storage, suedocument_name)
        self.document.save(suedocument) 
        
        return suedocument_name, suedocument

    def drive_upload(self, word_name, word_path):
        pdf_name = word_name
        pdf_path = self.doc_process.convert_word_to_pdf(word_name, word_path) # 워드 문서 pdf로 변환한 뒤, 로컬 pdf 경로 획득
        
        edit_link = self.doc_process.word_upload(word_name, word_path) # 워드 드라이브에 업로드 한 뒤 편집 링크 획득
        pdf_id = self.doc_process.pdf_upload(pdf_name, pdf_path) # pdf 드라이브에 업로드 한 뒤 pdf_id 획득
        
        with open(self.input_dic, "a", encoding="utf-8") as file:
            # 텍스트 추가
            file.write(word_name+"\n")
            file.write(str(self.form_data)+"\n")

        print(edit_link)
        print(pdf_id)
        print(word_name)
        
sue_gpt = sue_gpt(form_data)
sue_gpt.main()