from datetime import datetime
from docx import Document
from app.libs.form import storage_path
import os

def add_sentence(document: Document, target_sentence: str, new_text: str):
    for i, paragraph in enumerate(document.paragraphs):
        if target_sentence in paragraph.text:
            new_paragraph = document.add_paragraph(new_text)
            paragraph._element.addnext(new_paragraph._element)
            break

def format_korean_date(date_str: str) -> str:
    if not date_str:
        return ""
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            dt = datetime.strptime(date_str, fmt)
            return f"{dt.year}년 {dt.month}월 {dt.day}일"
        except ValueError:
            continue
    # 파싱 실패 시 원본 반환
    return date_str

def word_generate(llm_data, form_data):
    document = Document(os.path.join(os.path.dirname(__file__), '..', 'forms', '고소장 표준 양식(경찰청).docx'))

    now = datetime.now()
    
    table_requester = document.tables[0]
    for row in table_requester.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text == "고소인 성명":
                    para.text = form_data["고소인"]
                elif para.text == "고소인 주민등록번호":
                    para.text = form_data["고소인 주민등록번호"]
                elif para.text == "고소인 주소":
                    para.text = form_data["고소인 주소"]
                elif para.text == "고소인 직업":
                    para.text = form_data["고소인 직업"]
                elif para.text == "고소인 전화":
                    para.text = form_data["고소인 전화"]
                elif para.text == "고소인 이메일":
                    para.text = form_data["고소인 이메일"]
    table_receiver = document.tables[1]

    for row in table_receiver.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text == "피고소인 성명":
                    para.text = form_data["피고소인"]
                elif para.text == "피고소인 주민등록번호":
                    para.text = form_data["피고소인 주민등록번호"]
                elif para.text == "피고소인 주소":
                    para.text = form_data["피고소인 주소"]
                elif para.text == "피고소인 직업":
                    para.text = form_data["피고소인 직업"]
                elif para.text == "피고소인 전화":
                    para.text = form_data["피고소인 전화"]
                elif para.text == "피고소인 이메일":
                    para.text = form_data["피고소인 이메일"]
                elif para.text == "피고소인 기타사항":
                    para.text = llm_data["1"]

    add_sentence(document, "(죄명 및 피고소인에 대한 처벌의사 기재)", "\n"+llm_data["2"]) # \n 무조건 추가 필요
    add_sentence(document, "4. 범죄사실*", "\n"+llm_data["3"]) # \n 무조건 추가 필요
    add_sentence(document, "5. 고소이유", "\n"+llm_data["4"]) # \n 무조건 추가 필요
    add_sentence(document, "6. 증거자료", "\n"+llm_data["5"]) # \n 무조건 추가 필요

    checklist_table = document.tables[2]

    for row in checklist_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text == "본 고소장과 같은 내용의 고소장을 다른 검찰청 또는 경찰서에 제출하거나 제출하였던 사실이 ":
                    para.add_run(form_data["동일 건 고소 및 취소 사실 여부"])
                elif para.text == "본 고소장에 기재된 범죄사실과 관련된 사건 또는 공범에 대하여 검찰청이나 경찰서에서 수사 중에 ":
                    para.add_run(form_data["관련 형사사건 수사 유무"])

    formatted_date = format_korean_date(form_data["고소일자"])
    add_sentence(document, "8. 기타", "\n"+llm_data["6"]) # \n 무조건 추가 필요
    add_sentence(document, "무고죄로 처벌받을 것임을 서약합니다.", "\n                                                 "+formatted_date)
    add_sentence(document, "고소대리의 경우에는 제출인을 기재하여야 합니다.", "\n\n                                                   "+f"{form_data["제출 경찰서"]} 귀중")
    suedocument_name = "AI 고소장_"+str(form_data["고소인"])+"_"+str(form_data["고소 죄명"])+"_"+str(now.strftime("%H%M%S"))+".docx"
    suedocument_path = storage_path / suedocument_name
    document.save(suedocument_path) 
    
    return suedocument_path